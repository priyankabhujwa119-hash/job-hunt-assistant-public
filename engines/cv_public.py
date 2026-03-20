import io
import os
from docx import Document
from docx.shared import Pt
from datetime import date


def extract_cv_summary(cv_bytes):
    try:
        doc = Document(io.BytesIO(cv_bytes))
        full_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        summary_lines = []
        capture = False
        for line in full_text:
            if (
                "EXECUTIVE PROFILE" in line.upper()
                or "PROFILE" in line.upper()
                or "SUMMARY" in line.upper()
            ):
                capture = True
                continue
            if capture and line:
                if any(
                    s in line.upper()
                    for s in ["EXPERIENCE", "COMPETENCIES", "EDUCATION", "SKILLS", "CORE"]
                ):
                    break
                summary_lines.append(line)
            if len(summary_lines) >= 4:
                break
        return " ".join(summary_lines)
    except Exception as e:
        return ""


def tailor_cv_summary(cv_summary, job, profile, groq_key):
    if not groq_key or groq_key == "test_mode":
        return (
            f"Experienced talent acquisition leader with {profile.get('years_experience',10)}+ "
            f"years in RPO/MSP environments, seeking {job['title']} role at {job['company']}. "
            "Strong European market expertise with proven track record in delivering high-volume "
            "recruitment programmes. Available to relocate and contribute immediately."
        )

    try:
        from groq import Groq

        client = Groq(api_key=groq_key)
        prompt = f"""Rewrite this CV summary for a job application. Return ONLY the paragraph. 

 CANDIDATE: {profile.get('name')}, {profile.get('years_experience')} years experience 
 MARKETS: {profile.get('experience_markets',[])} 
 TARGET JOB: {job['title']} at {job['company']} in {job['location']} 
 JD: {job.get('description','')[:600]} 
 
 CURRENT SUMMARY: {cv_summary} 
 
 Rules: 
 - 4-5 sentences only 
 - Mirror JD keywords naturally 
 - Include metrics if available 
 - End with relocation/availability statement 
 - No client company names 
 - Return ONLY the paragraph"""

        r = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
        )
        return r.choices[0].message.content.strip()
    except Exception as e:
        return cv_summary


def generate_cover_letter(job, profile, groq_key):
    if not groq_key or groq_key == "test_mode":
        return f"""Dear Hiring Team, 

 Having followed {job['company']}'s work closely, I am excited to apply for the {job['title']} position. Your organisation's approach to talent acquisition aligns perfectly with my professional values and expertise. 
 
 With {profile.get('years_experience',10)}+ years leading recruitment functions across European, Indian, and Middle Eastern markets, I bring a proven track record in RPO delivery and strategic talent acquisition. I have consistently delivered results including 100% client retention and significant productivity improvements through AI-enabled processes. 
 
 I am fully committed to this opportunity and available to discuss my application at your earliest convenience. I look forward to contributing to {job['company']}'s continued success. 
 
 Best regards, 
 {profile.get('name','')} 
 {profile.get('email','')} 
 {profile.get('phone','')}"""

    try:
        from groq import Groq

        client = Groq(api_key=groq_key)
        prompt = f"""Write a 3-paragraph cover letter for {profile.get('name')} applying to {job['title']} at {job['company']}. 

 Profile: {profile.get('years_experience')} years, markets: {profile.get('experience_markets',[])}, relocate: {profile.get('relocate',True)} 
 JD: {job.get('description','')[:500]} 
 
 Rules: 
 - Para 1: Why this company and role 
 - Para 2: Why candidate fits with real metrics 
 - Para 3: Relocation commitment and close 
 - No client names 
 - Professional warm tone 
 - Return only letter body, no greeting/sign-off"""

        r = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
        )
        body = r.choices[0].message.content.strip()
        return f"""Dear Hiring Team, 

 {body} 
 
 Best regards, 
 {profile.get('name','')} 
 {profile.get('email','')} 
 {profile.get('phone','')}"""
    except Exception as e:
        return "Cover letter generation failed. Please try again."


def create_tailored_cv_bytes(cv_bytes, tailored_summary, job):
    try:
        doc = Document(io.BytesIO(cv_bytes))
        profile_found = False
        replaced = False
        for para in doc.paragraphs:
            if (
                "EXECUTIVE PROFILE" in para.text.upper()
                or "PROFILE" in para.text.upper()
                or "SUMMARY" in para.text.upper()
            ):
                profile_found = True
                continue
            if profile_found and not replaced and para.text.strip():
                if any(
                    s in para.text.upper()
                    for s in ["EXPERIENCE", "COMPETENCIES", "EDUCATION", "SKILLS", "CORE"]
                ):
                    break
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = tailored_summary
                else:
                    para.add_run(tailored_summary)
                replaced = True
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()
    except Exception as e:
        return cv_bytes


if __name__ == "__main__":
    print("CV public engine ready ✓")

